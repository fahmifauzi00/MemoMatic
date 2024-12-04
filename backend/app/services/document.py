# backend/app/services/document.py
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_docx(minutes_data: dict) -> str:
    doc = Document()
    
    # Add title
    title = doc.add_heading(minutes_data.get('title', 'Meeting Minutes'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary section
    if 'summary' in minutes_data:
        doc.add_heading('Summary', 1)
        doc.add_paragraph(minutes_data['summary'])
    
    # Add key points
    if 'key_points' in minutes_data and minutes_data['key_points']:
        doc.add_heading('Key Points', 1)
        for point in minutes_data['key_points']:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(point)
    
    # Add action items
    if 'action_items' in minutes_data and minutes_data['action_items']:
        doc.add_heading('Action Items', 1)
        for item in minutes_data['action_items']:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(item)
    
    # Add decisions
    if 'decisions' in minutes_data and minutes_data['decisions']:
        doc.add_heading('Decisions', 1)
        for decision in minutes_data['decisions']:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(decision)
    
    # Ensure temp directory exists
    PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".."))
    TEMP_DIR = os.path.join(PROJECT_ROOT, "temp")
    
    # Save document
    output_path = os.path.join(TEMP_DIR, 'minutes.docx')
    doc.save(output_path)
    return output_path